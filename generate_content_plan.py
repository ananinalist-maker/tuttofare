from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

BRAND_BLACK = RGBColor(0x0d, 0x0d, 0x0d)
BRAND_LIME = RGBColor(0x6d, 0x9b, 0x08)
ACCENT_BLUE = RGBColor(0x1a, 0x56, 0x76)
ACCENT_GREEN = RGBColor(0x2e, 0x7d, 0x32)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xff, 0xff, 0xff)
RED = RGBColor(0xc6, 0x28, 0x28)
ORANGE = RGBColor(0xe6, 0x5c, 0x00)

def shade(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = BRAND_BLACK
        r.font.name = 'Calibri'
    return h

def body(text, bold=False, size=None, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'; r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(10)
    return p

def table(headers, rows, col_widths=None, hdr_color="1a5676"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, hdr_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9); r.font.name = 'Calibri'
            if ri % 2 == 1: shade(c, "f0f0f0")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def content_block(title, format_type, funnel_stage, caption, notes=""):
    """Create a styled content post block"""
    p = doc.add_paragraph()
    r = p.add_run(f'{title}')
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT_BLUE

    t = doc.add_table(rows=3 if not notes else 4, cols=2)
    t.style = 'Table Grid'

    labels = ['Формат', 'Этап воронки', 'Текст поста']
    values = [format_type, funnel_stage, caption]
    if notes:
        labels.append('Примечания')
        values.append(notes)

    for i, (label, value) in enumerate(zip(labels, values)):
        lc = t.rows[i].cells[0]; lc.text = ''
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True; lr.font.size = Pt(9); lr.font.name = 'Calibri'
        shade(lc, "e8f0fe")
        lc.width = Cm(3)

        vc = t.rows[i].cells[1]; vc.text = ''
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(9); vr.font.name = 'Calibri'
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('VOLTIS FOOTBALL\nDEVELOPMENT CAMP')
r.font.size = Pt(36); r.font.color.rgb = BRAND_BLACK; r.bold = True; r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('КОНТЕНТ-ПЛАН')
r.font.size = Pt(24); r.font.color.rgb = ACCENT_BLUE; r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Анализ · Золотые куски · Все форматы\n8 недель готового контента для @voltis.camp')
r.font.size = Pt(14); r.font.color.rgb = MEDIUM_GRAY; r.font.name = 'Calibri'

for _ in range(7):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Июнь 2026 · Конфиденциально')
r.font.size = Pt(11); r.font.color.rgb = MEDIUM_GRAY; r.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════
# CONTENTS
# ═══════════════════════════════════════════════════

heading('СОДЕРЖАНИЕ', 1)

toc = [
    ('1.', 'Анализ: откуда берём контент'),
    ('2.', 'Золотые куски: 15 убойных тезисов'),
    ('3.', 'Контент-стратегия: столпы и тон'),
    ('4.', 'Reels: 10 готовых скриптов'),
    ('5.', 'Carousels: 8 готовых раскладок'),
    ('6.', 'Stories: шаблоны на каждый день'),
    ('7.', 'Посты: 12 готовых подписей'),
    ('8.', 'Рекламные креативы: тексты для Meta Ads'),
    ('9.', 'Контент-календарь на 8 недель'),
    ('10.', 'Хэштеги и визуальные правила'),
]
for num, title in toc:
    p = doc.add_paragraph()
    rn = p.add_run(f'{num}  '); rn.bold = True; rn.font.size = Pt(12); rn.font.color.rgb = ACCENT_BLUE; rn.font.name = 'Calibri'
    rt = p.add_run(title); rt.font.size = Pt(12); rt.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. ANALYSIS
# ═══════════════════════════════════════════════════

heading('1. АНАЛИЗ: ОТКУДА БЕРЁМ КОНТЕНТ', 1)

body('1.1 Источники контента из проекта', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Источник', 'Что даёт для контента', 'Сколько единиц'],
    [
        ['4 рекламных оффера (EN + RU)', 'Готовые продающие тексты, протестированные', '4 поста + 4 рекламы'],
        ['2 видео-скрипта (суфлер)', 'Готовые Reels-скрипты', '2 Reels + нарезка'],
        ['Конкурентный анализ (12 академий)', 'Сравнения, факты о рынке, USP', '3 carousel + 5 постов'],
        ['Аналитика лидов (505 шт.)', 'Реальные вопросы и возражения родителей', 'FAQ-серия, Stories'],
        ['2 покупателя (Julian SE, Jarno BE)', 'Отзывы, истории, кейсы', '4 поста + 2 Reels'],
        ['Продукт: Player Report, 10+ метрик', 'Экспертный контент, образование', '5 carousel + 3 Reels'],
        ['Тренер Real Madrid', 'Авторитет, экспертность', '3 Reels + 2 поста'],
        ['Локация: Кастельон, 4★ отель, побережье', 'Визуальный/lifestyle контент', '5 Reels + Stories'],
        ['Правки Adrian (13.05)', 'Tone of voice: без "мечты", без "станет профи"', 'Фильтр для всего контента'],
    ],
    col_widths=[5, 6, 4]
)

body('1.2 Что есть и чего не хватает', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Есть', 'Нет (нужно создать)'],
    [
        ['Тексты офферов (EN)', 'Фото/видео с реального лагеря'],
        ['Видео-скрипты', 'UGC от покупателей (Julian, Jarno)'],
        ['Конкурентный анализ', 'Контент от тренера (talking head)'],
        ['PDF-презентация (18 стр.)', 'Behind-the-scenes материалы'],
        ['Брендовые цвета (#0d0d0d + #b5f03c)', 'Фото площадки, отеля, побережья'],
        ['Instagram @voltis.camp', 'Шаблоны для Stories/Carousel в Canva'],
    ],
    col_widths=[8, 8]
)

body('1.3 Tone of Voice (согласован с Adrian)', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Можно', 'Нельзя'],
    [
        ['Professional development', '"Мечта" / "dream"'],
        ['"Level up", "next step"', '"Станет профессионалом"'],
        ['"European club benchmarks"', '"Откроются двери"'],
        ['"Outgrown their current level"', '"Застряли" / "stuck"'],
        ['CAMP / PROGRAM крупно', '"Сборная" — фокус на клубы'],
        ['Факты, метрики, конкретика', 'Пустые обещания'],
    ],
    col_widths=[8, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 2. GOLDEN NUGGETS
# ═══════════════════════════════════════════════════

heading('2. ЗОЛОТЫЕ КУСКИ: 15 УБОЙНЫХ ТЕЗИСОВ', 1)

body('Это — ядро всего контента. Каждый тезис = минимум 3 единицы контента (Reel + Carousel + пост). Используй их как заголовки, хуки, первые строки.', size=10, color=MEDIUM_GRAY)

nuggets = [
    ('ПРОДУКТ', [
        ('€357/день', 'За эти деньги: 2 тренировки с тренером Real Madrid, 4★ отель, all inclusive, план развития на 6 месяцев. Годовая программа в академии — от €18,000.'),
        ('1 тренер на 3 игрока', 'У большинства лагерей — 1:10 или 1:15. В Voltis — персональное внимание к каждому.'),
        ('Player Report 8–10 страниц', 'Ни один из 12 конкурентов не даёт письменный отчёт + план развития + 30 дней поддержки после. Все продают опыт. Voltis продаёт ответ и план.'),
        ('10+ метрик тестирования', 'First touch, decision speed, body orientation, passing accuracy, 1v1... — то же, что оценивают скауты европейских клубов.'),
        ('30 дней поддержки после лагеря', 'Лагерь не заканчивается на 7-й день. Тренер остаётся на связи и корректирует план.'),
    ]),
    ('РЫНОК', [
        ('0 из 12 конкурентов дают то же самое', 'Мы проанализировали Villarreal, GIMFA, Alicante FA, Siello, Ertheo и ещё 7. Никто не даёт: отчёт + план + поддержку + видеоанализ + 10 метрик.'),
        ('€2,500 vs €18,000–50,000/год', 'Voltis за 7 дней даёт то, что годовые программы дают за €18K–50K — и добавляет то, чего у них нет (отчёт, план, поддержка).'),
        ('Скандинавия и Бенилюкс — самый горячий отклик', 'Норвегия, Бельгия, Нидерланды, Швеция, Дания — эти страны лучше всего реагируют на предложение.'),
    ]),
    ('SOCIAL PROOF', [
        ('Julian из Швеции', '"My son improved more in 7 days than in 6 months at his local club." — 2 позиции: 12 и 16 лет.'),
        ('Jarno из Бельгии', 'Оплатил полностью €2,500 без торга. Нашёл через рекламу, прошёл форму с вопросом о бюджете — и сразу купил.'),
        ('Tracy из Бельгии привёл друга', 'Ещё до лагеря рекомендовал другой семье. Реферальный канал работает органически.'),
    ]),
    ('ВОЗРАЖЕНИЯ (из 505 лидов)', [
        ('"Дорого" — 34% лидов', 'Самое частое возражение. Контент должен переупаковать цену: €357/день, сравнение с академиями, рассрочка.'),
        ('"Лето дорого → осень"', 'Многие лиды отложили на осень. Это НЕ отказ — это отложенный спрос. Нужен контент-прогрев до осени.'),
        ('"Ребёнок не готов"', 'Это НЕ просмотр и НЕ отбор. Это тренировка + оценка для ЛЮБОГО уровня. Контент должен это объяснять.'),
    ]),
]

for category, items in nuggets:
    body(f'  {category}', bold=True, size=11, color=ACCENT_BLUE)
    for i, (hook, detail) in enumerate(items, 1):
        p = doc.add_paragraph()
        rn = p.add_run(f'  {hook}  ')
        rn.bold = True; rn.font.size = Pt(10); rn.font.name = 'Calibri'; rn.font.color.rgb = BRAND_BLACK
        rd = p.add_run(f'— {detail}')
        rd.font.size = Pt(9); rd.font.name = 'Calibri'; rd.font.color.rgb = MEDIUM_GRAY
        p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 3. CONTENT STRATEGY
# ═══════════════════════════════════════════════════

heading('3. КОНТЕНТ-СТРАТЕГИЯ: СТОЛПЫ И ТОН', 1)

body('3.1 Четыре контент-столпа', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Столп', 'Доля', 'Цель', 'Примеры'],
    [
        ['EDUCATE\n(Образование)', '30%', 'Показать экспертность,\nдать ценность бесплатно', 'Как скауты оценивают, 5 ошибок игроков,\nчто такое Player Report, 10 метрик'],
        ['PROVE\n(Доказательства)', '25%', 'Убрать сомнения,\nпоказать результат', 'Отзывы Julian/Jarno, сравнение с конкурентами,\nпример Player Report, цифры'],
        ['SHOW\n(Атмосфера)', '25%', 'Вызвать желание,\nпоказать среду', 'День в лагере, площадка, отель, тренер,\nпобережье, тренировки, еда'],
        ['SELL\n(Продажа)', '20%', 'Конвертировать в лид,\nподтолкнуть к действию', 'Early Bird, осталось X мест, CTA,\nсемейный пакет, FAQ'],
    ],
    col_widths=[3, 1.5, 3.5, 8]
)

body('3.2 Правило: контент ≠ реклама', bold=True, size=12, color=RED)

p = doc.add_paragraph()
r = p.add_run('80% контента даёт ценность. 20% продаёт.')
r.bold = True; r.font.size = Pt(11); r.font.name = 'Calibri'

p = doc.add_paragraph()
r = p.add_run('Ошибка текущего подхода: 100% контента — "запишись на лагерь". Холодному человеку нет дела до лагеря. Ему интересно: как оценивают игроков? какие ошибки делает его ребёнок? что такое испанская методология? Сначала ценность → потом доверие → потом продажа.')
r.font.size = Pt(10); r.font.name = 'Calibri'

body('3.3 Визуальный стиль', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Элемент', 'Стандарт'],
    [
        ['Основные цвета', 'Чёрный #0d0d0d + лаймовый #b5f03c'],
        ['Шрифты', 'Жирный sans-serif для заголовков, чистый для текста'],
        ['Фото-стиль', 'Тёмный, кинематографичный, ночной стадион + золотой час'],
        ['Альтернативный стиль', 'Яркий Mediterranean lifestyle (как Le Court Tennis)'],
        ['Слова CAMP / PROGRAM', 'ВСЕГДА крупно, на уровне заголовка (требование Adrian)'],
        ['Формат постов', '1:1 (1080x1080) для ленты, 9:16 для Reels/Stories'],
        ['Субтитры в видео', 'ОБЯЗАТЕЛЬНО — 80% смотрят без звука'],
    ],
    col_widths=[4, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 4. REELS SCRIPTS
# ═══════════════════════════════════════════════════

heading('4. REELS: 10 ГОТОВЫХ СКРИПТОВ', 1)

body('Правила: формат 9:16, длина 15–60 сек, субтитры обязательно, первые 3 секунды = хук, финал = CTA на чёрном фоне.', size=9, color=MEDIUM_GRAY)

reels = [
    ('REEL 1 — "How scouts actually evaluate your child"', 'EDUCATE', '45–60 сек',
     '''Hook (3 сек): "Most parents don't know how European scouts actually rate players."

Body: "They don't care about goals scored. They look at 10 specific things:"
1. First touch quality
2. Decision speed under pressure
3. Body orientation when receiving
4. Passing accuracy — short AND long
5. 1v1 — attacking and defending
6. Movement off the ball
7. Game intelligence
8. Composure under press
9. Tactical awareness
10. Physical data

"At Voltis Camp, we test ALL of these. Your child leaves with a written report — exactly where they stand vs. European club standards."

CTA: "Link in bio."
NO SELLING — only education.''',
     'Экран: тренер на поле, вечер, прожекторы. Каждая метрика = 1 карточка субтитров. Графика поверх видео.'),

    ('REEL 2 — "5 mistakes every young player makes"', 'EDUCATE', '30–45 сек',
     '''Hook: "I trained in the Real Madrid system. Here are 5 mistakes I see in every young player."

1. "Looking down when receiving the ball"
2. "Always choosing the safe pass"
3. "No movement before the ball arrives"
4. "Panicking under pressure"
5. "Training without a plan"

"The good news? All 5 are fixable. Some in just 7 days."

CTA: "Follow for more football development tips."''',
     'Тренер говорит на камеру (talking head). Каждая ошибка = смена ракурса. Динамичный монтаж.'),

    ('REEL 3 — "Day in the life at Voltis Camp"', 'SHOW', '45–60 сек',
     '''Hook: "What does a professional football training day look like for a 14-year-old?"

"7:30 — Wake up. 4-star hotel on the Spanish coast."
"8:00 — Breakfast designed by a sports nutritionist."
"9:00 — Morning session. Technical drills. Spanish methodology."
"12:00 — Lunch + recovery."
"15:00 — Video analysis. Where you went wrong. How to fix it."
"17:00 — Evening session. Tactical work. Game situations."
"20:00 — Dinner. Player Report update."
"This is not a summer camp. This is professional development."

CTA: "Autumn camp — link in bio."''',
     'B-roll стиль: если нет реального видео — AI-визуализация + фото площадки/отеля. Тёплые тона, золотой час.'),

    ('REEL 4 — "€357 per day. Here\'s what you get."', 'PROVE', '30 сек',
     '''Hook: "People say €2,500 for a football camp is expensive. Let me break it down."

"€357 per day. That's:"
"2 training sessions with a Real Madrid system coach"
"4-star hotel on the Mediterranean"
"3 meals a day from a sports nutritionist"
"Video analysis of every session"
"A personal Player Report — 8 to 10 pages"
"A 6-month development plan"
"30 days of post-camp coaching support"

"A year at a Spanish football academy? €18,000 to €50,000."

CTA: "Still think it's expensive? Link in bio."''',
     'Текст поверх B-roll. Каждая строка = быстрая смена кадра. Финал — цифра €18,000–50,000 крупно.'),

    ('REEL 5 — "What parents don\'t know about football development"', 'EDUCATE', '30 сек',
     '''Hook: "Your child trains 3 times a week. Plays matches on weekends. But here's what's missing."

"No professional evaluation."
"No written development plan."
"No correction from someone who trained at the top level."
"No benchmark against European standards."
"Your child might be talented. But without these — they're training blind."

CTA: "Follow @voltis.camp — we fix this in 7 days."''',
     'Тёмный фон, текст появляется строка за строкой. Минимальное видео, максимум типографика.'),

    ('REEL 6 — Julian\'s story (testimonial)', 'PROVE', '30 сек',
     '''Hook: "Julian from Sweden came to Voltis Camp. This is what happened."

"Before: A talented player with no clear direction."
"After 7 days:"
"— Personal evaluation against EU club benchmarks"
"— Written Player Report: strengths, weaknesses, exact numbers"
"— 6-month plan his local coach now follows"

His dad: "My son improved more in 7 days than in 6 months at his local club."

CTA: "Autumn camp — limited spots. Link in bio."''',
     'Идеально: реальное видео Julian + цитата отца. Если нет — текст поверх фото + голос за кадром.'),

    ('REEL 7 — "Real Madrid coaching system — what does it mean?"', 'EDUCATE', '30 сек',
     '''Hook: "We say 'Real Madrid coaching system.' But what does that actually mean?"

"It means:"
"Technical precision over physical power."
"Decision-making speed — not just running fast."
"Building from the back. Positional play."
"Rondos. Tactical awareness. Game intelligence."

"This is what separates European academies from local clubs."
"And this is what your child gets in 7 days at Voltis."

CTA: "Details in bio."''',
     'Тренер на камеру, на поле. Уверенно, спокойно, с экспертизой.'),

    ('REEL 8 — "Is your child ready?" (engagement)', 'EDUCATE', '15–20 сек',
     '''Hook: "Quick test. Is your child ready for professional football development?"

"Can they receive the ball without looking down?"
"Can they make a decision in under 2 seconds?"
"Do they know their weaknesses — specifically?"
"Do they have a written development plan?"

"If you answered 'no' to any of these — we can help."

CTA: "Follow for more. Link in bio."''',
     'Быстрый формат, текст на экране, каждый вопрос = новый кадр.'),

    ('REEL 9 — Location reveal', 'SHOW', '20–30 сек',
     '''Hook: "This is where your child will train."

"Castellón, Spain. Mediterranean coast."
"Professional football pitch."
"4-star hotel — 5 minutes from the sea."
"Sports nutritionist. 3 meals daily."
"Evening sessions under floodlights."

"7 days. Real training. Real results."

CTA: "Autumn camp — reserve your spot."''',
     'Дрон-шоты / фото площадки, отеля, моря. Красивая нарезка. Музыка — эпик/кинематограф.'),

    ('REEL 10 — "This is NOT a summer camp"', 'SELL', '20 сек',
     '''Hook: "Let me be clear. This is NOT a summer camp."

"No arts and crafts."
"No campfire songs."
"No participation trophies."

"This is:"
"Professional football evaluation."
"Technical correction by a Real Madrid system coach."
"A personal development plan for the next 6 months."

"Voltis Development Camp. Spain."

CTA: "Apply now — link in bio."''',
     'Жёсткий, контрастный монтаж. Чёрный фон → яркие кадры тренировок. Текст крупно.'),
]

for title, pillar, duration, script, notes in reels:
    content_block(
        title,
        f'Reel 9:16 · {duration}',
        f'Столп: {pillar}',
        script,
        notes
    )

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 5. CAROUSELS
# ═══════════════════════════════════════════════════

heading('5. CAROUSELS: 8 ГОТОВЫХ РАСКЛАДОК', 1)

body('Формат: 1:1 (1080x1080), 5–10 слайдов, свайп-стопперы, последний слайд = CTA.', size=9, color=MEDIUM_GRAY)

carousels = [
    ('CAROUSEL 1 — "10 метрик, по которым скауты оценивают игрока"', 'EDUCATE',
     '''Слайд 1 (обложка): "10 METRICS EUROPEAN SCOUTS USE TO EVALUATE YOUR CHILD"
Подзаголовок: "Most parents have no idea. Swipe →"

Слайд 2: "#1 — First Touch Quality"
Пояснение: "How clean is the first control? Can they receive under pressure?"

Слайд 3: "#2 — Decision Speed"
"How fast do they choose: pass, dribble, or shoot?"

Слайд 4: "#3 — Body Orientation"
"Where is their body facing when they receive? This shows game awareness."

Слайд 5: "#4 — Passing Accuracy"
"Short AND long. Both feet. Under pressure."

Слайд 6: "#5 — 1v1 Ability"
"Attacking: can they beat a defender? Defending: can they stop one?"

Слайд 7: "#6–10 — Movement, IQ, Composure, Tactics, Physicals"
"The full picture. Not just 'is he fast?'"

Слайд 8 (CTA): "At Voltis Camp, we test ALL 10."
"Your child leaves with a written Player Report."
"Autumn camp — link in bio."''',
     'Дизайн: тёмный фон, лаймовые акценты, номера крупно. Можно использовать иконки для каждой метрики.'),

    ('CAROUSEL 2 — "Voltis vs. Академии: сравнение"', 'PROVE',
     '''Слайд 1: "VOLTIS CAMP vs. SPANISH FOOTBALL ACADEMIES — honest comparison"

Слайд 2: "Duration"
Voltis: 7 days | Academies: 10–12 months
"You don't need a year to know where your child stands."

Слайд 3: "Price"
Voltis: €2,500 | Academies: €18,000–50,000
"Same coaches. Same methodology. 7–20x less."

Слайд 4: "Player Report"
Voltis: ✓ 8–10 pages | Academies: ✗ Usually none
"We give written feedback. Most academies don't."

Слайд 5: "Development Plan"
Voltis: ✓ 6-month plan | Academies: ✗ No takeaway plan
"Your child leaves with a roadmap."

Слайд 6: "Coach:Player Ratio"
Voltis: 1:3 | Academies: 1:10–15
"3x more attention per player."

Слайд 7: "Post-Camp Support"
Voltis: ✓ 30 days | Academies: ✗ Zero
"The camp doesn't end on day 7."

Слайд 8 (CTA): "All of this for €357/day. All inclusive."
"Autumn camp — reserve your spot."''',
     'Формат split-screen: левая колонка Voltis (лаймовый), правая — Академии (серый).'),

    ('CAROUSEL 3 — "Что входит в 7 дней" (раскладка по дням)', 'SHOW',
     '''Слайд 1: "7 DAYS AT VOLTIS CAMP — what actually happens"

Слайд 2: "DAY 1 — Assessment"
"Initial testing. 10+ metrics. Baseline report."

Слайд 3: "DAY 2–3 — Technical Correction"
"Spanish methodology. Fix the key mistakes. Individual focus."

Слайд 4: "DAY 4–5 — Tactical Development"
"Game situations. Decision making. Positional play."

Слайд 5: "DAY 6 — Competition Day"
"Real match. Coaches observe. Video analysis."

Слайд 6: "DAY 7 — Final Report"
"Player Report delivered. 6-month plan. Coach Q&A."

Слайд 7: "AFTER — 30 Days Support"
"Coach stays in touch. Plan adjustments. Progress check."

Слайд 8 (CTA): "More than a camp. A development program."
"Link in bio."''',
     'Визуал: таймлайн, каждый день = новый цвет/иконка. Чистый, информативный дизайн.'),

    ('CAROUSEL 4 — "FAQ — ваши вопросы"', 'SELL',
     '''Слайд 1: "YOUR QUESTIONS ABOUT VOLTIS CAMP — answered honestly"

Слайд 2: "Is €2,500 worth it for 7 days?"
"€357/day: professional coaching, 4★ hotel, all meals, personal report, 6-month plan. A year at a Spanish academy: €18K–50K."

Слайд 3: "My child isn't good enough."
"We take ages 10–19 at ALL levels. This is NOT a tryout — it's training + evaluation. Your child gets a plan tailored to THEIR level."

Слайд 4: "Is it safe?"
"1:3 coach-to-player ratio. Medical staff on site. 4★ hotel with 24/7 supervision. All inclusive."

Слайд 5: "What about travel?"
"Fly to Valencia. We handle the rest: transfer, hotel, meals, training. Family package available — you vacation in Spain while your child trains."

Слайд 6: "What does my child take home?"
"Player Report (8–10 pages), video analysis, 6-month development plan, 30 days of coach support."

Слайд 7 (CTA): "More questions? DM us or WhatsApp — link in bio."''',
     'Формат: вопрос крупно сверху (белый на тёмном), ответ — текст ниже.'),

    ('CAROUSEL 5 — "Player Report — что внутри"', 'PROVE',
     '''Слайд 1: "PLAYER REPORT — what your child actually receives"
"8–10 pages of professional evaluation"

Слайд 2: "Page 1 — Player Profile"
"Name, age, position, physical data, training history"

Слайд 3: "Pages 2–4 — Technical Assessment"
"10+ metrics scored against European club benchmarks"

Слайд 4: "Page 5 — Tactical Analysis"
"Positioning, decision-making, game intelligence"

Слайд 5: "Pages 6–7 — Video Breakdown"
"Key moments from training sessions — what went right, what to fix"

Слайд 6: "Pages 8–10 — Development Plan"
"6-month roadmap: specific drills, goals, milestones"

Слайд 7 (CTA): "No other camp gives this."
"We checked 12 competitors. Zero offer a written report + plan + post-camp support."
"Link in bio."''',
     'Если есть реальные скриншоты Player Report (замазать данные) — идеально. Если нет — мокап.'),

    ('CAROUSEL 6 — "5 причин почему местный клуб — не достаточно"', 'EDUCATE',
     '''Слайд 1: "5 REASONS YOUR LOCAL CLUB ISN'T ENOUGH"
"(And what to do about it)"

Слайд 2: "#1 — No individual evaluation"
"Group training doesn't tell you where YOUR child stands."

Слайд 3: "#2 — No European benchmark"
"Your child might be the best in the team — but is that enough for a European club?"

Слайд 4: "#3 — No development plan"
"Training 3x/week without a plan is training blind."

Слайд 5: "#4 — Same methodology year after year"
"Most local coaches haven't trained in a top European system."

Слайд 6: "#5 — No honest feedback"
"Local coaches don't want to upset parents. We give you the truth — with a plan to fix it."

Слайд 7 (CTA): "Voltis Camp: 7 days to get the answers local clubs can't give."
"Link in bio."''',
     'Провокационный, но не оскорбительный. Тон: "мы не против местных клубов — мы дополняем".'),

    ('CAROUSEL 7 — "Что говорят родители" (отзывы)', 'PROVE',
     '''Слайд 1: "WHAT PARENTS ARE SAYING ABOUT VOLTIS CAMP"

Слайд 2: Julian's dad (Sweden)
"My son improved more in 7 days than in 6 months at his local club."
Фото/аватар + флаг 🇸🇪

Слайд 3: Jarno's family (Belgium)
"We knew this was the right investment from the first conversation."
Бюджет: €2,500 — оплатил полностью, без торга.
Фото/аватар + флаг 🇧🇪

Слайд 4: Tracy (Belgium)
"I recommended Voltis to another family before even going myself."
Флаг 🇧🇪 + пометка "word of mouth"

Слайд 5 (CTA): "Join them this autumn."
"Autumn camp — early bird pricing available."
"Link in bio."''',
     'Нужны реальные фото или хотя бы имена. Если нет фото — текстовые цитаты на тёмном фоне.'),

    ('CAROUSEL 8 — "Early Bird: 3 причины записаться сейчас"', 'SELL',
     '''Слайд 1: "3 REASONS TO BOOK VOLTIS AUTUMN CAMP NOW"

Слайд 2: "#1 — Save €300"
"Early Bird price: €2,200 instead of €2,500."
"Until July 31 only."

Слайд 3: "#2 — Only 15 spots total"
"1:3 coach ratio = maximum 15 players."
"First come, first served."

Слайд 4: "#3 — Family package: €1,300"
"Your child trains. You vacation in Spain."
"Mediterranean coast, October weather: perfect."

Слайд 5 (CTA): "Early Bird ends [date]."
"DM us or tap the link in bio to reserve."''',
     'Яркий, urgency-дизайн. Таймер / обратный отсчёт визуально. Лаймовые акценты.'),
]

for title, pillar, slides, notes in carousels:
    content_block(title, 'Carousel 1:1 · 5–8 слайдов', f'Столп: {pillar}', slides, notes)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 6. STORIES
# ═══════════════════════════════════════════════════

heading('6. STORIES: ШАБЛОНЫ НА КАЖДЫЙ ДЕНЬ', 1)

body('Публиковать 2–3 Stories в день. Микс форматов: опрос, вопрос, факт, behind the scenes. Stories — главный инструмент вовлечения и "прогрева" подписчиков.', size=10, color=MEDIUM_GRAY)

table(
    ['Тип', 'Пример', 'Частота', 'Цель'],
    [
        ['Опрос', '"At what age should a player start professional training?\n○ 8–10  ○ 11–13  ○ 14+"', '2 раза/нед', 'Engagement + data'],
        ['Вопрос (стикер)', '"Ask us anything about football development in Spain 🇪🇸"', '1 раз/нед', 'FAQ контент + вовлечение'],
        ['Факт дня', '"Did you know? The average LaLiga academy tests 10+ metrics before signing a player."', '3 раза/нед', 'Образование'],
        ['Обратный отсчёт', '"Autumn Camp — [X] days until Early Bird ends ⏰"', '1 раз/нед', 'Urgency'],
        ['Это или То', '"Better for development:\n🅰️ More matches  🅱️ More individual training"', '1 раз/нед', 'Engagement'],
        ['Behind the scenes', 'Фото/видео: площадка, отель, тренер готовит план, побережье', '2 раза/нед', 'Атмосфера + доверие'],
        ['Отзыв', 'Скриншот переписки / цитата Julian или Jarno', '1 раз/нед', 'Social proof'],
        ['CTA-story', '"DM us \'CAMP\' and we\'ll send you the full program 📩"', '1 раз/нед', 'Лиды из Instagram'],
        ['Репост Reel', 'Репост нового Reel в Stories + стикер "Смотрели?"', 'При публикации', 'Охват Reel'],
        ['Quiz', '"How many metrics do scouts evaluate? ○ 3  ○ 5  ○ 10+"', '1 раз/нед', 'Engagement + образование'],
    ],
    col_widths=[3, 6, 2, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 7. POSTS
# ═══════════════════════════════════════════════════

heading('7. ПОСТЫ: 12 ГОТОВЫХ ПОДПИСЕЙ', 1)

posts = [
    ('POST 1 — Вводный (о лагере)', 'SELL', 'Single image 1:1',
     '''Professional football development. Not a summer camp.

7 days in Spain. 2 training sessions every day.
Real Madrid coaching system. 1 coach per 3 players.

What your child takes home:
📋 Player Report — 8–10 pages of professional evaluation
📊 10+ metrics tested against European club standards
🎯 6-month development plan
📹 Video analysis of every session
📞 30 days of post-camp coach support

Ages 10–19 · All inclusive · Family package available.
Autumn camp — link in bio.

#footballcamp #footballdevelopment #youthfootball #footballspain'''),

    ('POST 2 — Образовательный (метрики)', 'EDUCATE', 'Carousel (см. Carousel 1)',
     '''When a scout watches your child play, they're not counting goals.

They evaluate 10 specific metrics. Swipe to see all 10 →

Most players never find out where they actually stand.
That's the gap between talent and opportunity.

At Voltis Camp, we close that gap in 7 days.

Save this post. Send it to a football parent who needs to see it.

#footballscout #youthfootball #footballdevelopment #footballmetrics'''),

    ('POST 3 — Social Proof (Julian)', 'PROVE', 'Single image / Video testimonial',
     '''"My son improved more in 7 days than in 6 months at his local club."

Julian came to Voltis Camp from Sweden 🇸🇪
Before: Talented but no clear development path.
After: Player Report with specific metrics. 6-month plan his local coach now follows.

Every parent wants the best for their child.
Most don't know what "the best" looks like in football development.

We show you — in writing. With numbers. With a plan.

Autumn camp opens in October. Early Bird pricing in bio.

#footballcamp #playerdevelopment #footballspain'''),

    ('POST 4 — Сравнение цен', 'PROVE', 'Carousel (см. Carousel 2)',
     '''Is €2,500 expensive for a football camp?

Let's do the math:
€357/day.

For that, your child gets:
→ 2 training sessions with a Real Madrid system coach
→ 4★ hotel on the Mediterranean
→ All meals from a sports nutritionist
→ Personal Player Report (8–10 pages)
→ 6-month development plan
→ 30 days of post-camp support

A year at a Spanish football academy: €18,000–50,000.
Voltis gives you the evaluation + plan for 7x–20x less.

Still think it's expensive?

#footballcost #footballacademy #voltiscamp'''),

    ('POST 5 — Тренер', 'EDUCATE', 'Single image / Video',
     '''Meet the coach.

Trained in the Real Madrid system.
Not just drills — a complete methodology:

→ Technical precision over brute force
→ Decision-making speed
→ Positional play and game intelligence
→ Building from the back

This is what your child gets at Voltis Camp.
Not a local coach with good intentions.
A professional who knows exactly what European clubs want to see.

1 coach per 3 players. Personal attention guaranteed.

#realmadrid #footballcoach #spanishfootball #footballmethodology'''),

    ('POST 6 — "Это не летний лагерь"', 'SELL', 'Single image (statement)',
     '''Let's be clear.

This is NOT a summer camp.
No campfire. No group games. No participation medals.

This is:
Professional football evaluation.
Technical correction.
Tactical development.
A written plan for the next 6 months.

7 days. Spain. All inclusive.
Your child leaves with answers — not just memories.

Voltis Development Camp.
Autumn spots opening. Link in bio.

#footballcamp #profootball #footballdevelopment'''),

    ('POST 7 — Behind the scenes (локация)', 'SHOW', 'Single image / Carousel',
     '''Castellón, Spain. Mediterranean coast.

Where your child will train:
☀️ Professional football pitch
🏨 4-star hotel — 5 min from the sea
🍽️ Meals designed by a sports nutritionist
📹 Video analysis room
⚽ Evening sessions under floodlights

October weather: 20–25°C. Perfect for football.
Perfect for family — family package available.

Autumn camp — details in bio.

#footballspain #castellon #footballcamp #mediterranean'''),

    ('POST 8 — Возражение "дорого"', 'PROVE', 'Single image (text post)',
     '''The most common question we get: "Isn't it too expensive?"

Here's what we say:

→ €357/day is less than most parents spend on a week of vacation
→ A year at Villarreal Academy: €25,000
→ A year at GIMFA Getafe: €49,000
→ Voltis: same methodology, fraction of the price

But the real question isn't cost — it's value.

What's the cost of your child training for years without a professional evaluation?
Without knowing where they actually stand?
Without a plan?

That's the expensive option.

#footballinvestment #footballparents #youthfootball'''),

    ('POST 9 — Early Bird оффер', 'SELL', 'Single image (offer)',
     '''📢 VOLTIS AUTUMN CAMP — Early Bird is OPEN

October 2026. Castellón, Spain. 7 days.

Early Bird price: €2,200 (save €300)
Family package: €1,300 (save €200)

What's included:
✅ 2 daily sessions · Real Madrid coaching system
✅ 4★ hotel · All inclusive
✅ Player Report · 10+ metrics
✅ 6-month development plan
✅ 30 days post-camp support

Only 15 spots. Early Bird ends July 31.

DM us "CAMP" or tap the link in bio.

#earlybird #footballcamp #voltiscamp #autumncamp'''),

    ('POST 10 — Player Report пример', 'PROVE', 'Carousel (см. Carousel 5)',
     '''This is what your child takes home from Voltis Camp.

Not a certificate. Not a photo.

A Player Report — 8 to 10 pages:
📊 10+ metrics scored against EU club standards
🎯 Strengths and weaknesses — with numbers
📹 Video breakdown of key moments
📋 6-month development plan — specific drills and goals

We checked 12 other football camps and academies.
ZERO of them give a written report + plan + post-camp support.

Swipe to see what's inside →

#playerreport #footballdevelopment #footballevaluation'''),

    ('POST 11 — Семейный пакет', 'SELL', 'Single image',
     '''Your child trains. You vacation.

Family Package — €1,500 (Early Bird: €1,300)

While your child is on the pitch with a Real Madrid system coach:
☀️ You're on the Mediterranean coast
🏨 4-star hotel
🍽️ All meals included
🏖️ Castellón beaches, old town, seafood

October in Spain: 20–25°C. No crowds. Perfect weather.

Combine football development with a family holiday.
DM us or link in bio.

#familyfootball #footballholiday #spaintravel #voltiscamp'''),

    ('POST 12 — Urgency / last spots', 'SELL', 'Single image',
     '''⚠️ VOLTIS AUTUMN CAMP — [X] SPOTS LEFT

We keep groups small for a reason:
1 coach per 3 players.
Maximum 15 participants.

That means when spots are gone — they're gone.

Early Bird (€2,200) ends [date].
After that: full price €2,500 — or waitlist.

7 days in Spain.
Professional evaluation. Development plan. Real coaching.

Don't wait until it's full.

DM "CAMP" or link in bio.

#limitedspots #footballcamp #voltiscamp #booknow'''),
]

for title, pillar, format_type, caption in posts:
    content_block(title, format_type, f'Столп: {pillar}', caption)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 8. AD CREATIVES
# ═══════════════════════════════════════════════════

heading('8. РЕКЛАМНЫЕ КРЕАТИВЫ: ТЕКСТЫ ДЛЯ META ADS', 1)

body('8.1 Контентная реклама (Video Views — Неделя 1–2)', bold=True, size=12, color=ACCENT_BLUE)

body('Цель: собрать аудиторию, НЕ продавать. Текст минимальный — видео делает работу.', size=10, color=MEDIUM_GRAY)

table(
    ['Ad', 'Primary Text', 'Headline', 'CTA Button'],
    [
        ['Video A\n"How scouts evaluate"',
         'Most parents don\'t know how European scouts actually rate players.\nWatch this — it might change how you see your child\'s training.',
         'The 10 Metrics That Matter',
         'Learn More'],
        ['Video B\n"5 mistakes"',
         'I trained in the Real Madrid system.\nHere are 5 mistakes I see in every young player — and how to fix them.',
         '5 Mistakes Young Players Make',
         'Watch More'],
        ['Video C\n"Day in the life"',
         'What does professional football training look like for a teenager?\n7:30 AM to 8 PM — every minute has a purpose.',
         'A Day at Voltis Camp',
         'Learn More'],
    ],
    col_widths=[2.5, 6, 3.5, 2]
)

body('8.2 Проспектинг реклама (Lead Gen — Неделя 3–8)', bold=True, size=12, color=ACCENT_BLUE)

body('Связка-победитель: статичные креативы + LEAD_GENERATION + EU list гео.', size=10, color=MEDIUM_GRAY)

table(
    ['Ad', 'Primary Text', 'Headline', 'CTA Button'],
    [
        ['Offer 1\n(лучший CTR 5.19%)',
         '⚽ Professional football camp in Spain for players aged 10–19.\n🗓 7 days · ⚡ 2 sessions daily · 👨‍🏫 Real Madrid system coach.\n📋 Get a personal development plan for reaching professional clubs.\n📍 Spain · ✅ All inclusive · 👨‍👩‍👦 Family package available.',
         'Professional Football Development Camp',
         'Sign Up'],
        ['Offer 3\n(первый шаг)',
         '🎯 Take the first step toward professional football.\n⚽ Camp in Spain — ages 10–19. In 7 days:\n💪 Level up skills\n✅ Fix key mistakes — Spanish methodology\n📊 Evaluate against European benchmarks\n📋 Get a 6-month plan\n👨‍🏫 Real Madrid system coach.',
         'Your Child\'s First Real Step',
         'Sign Up'],
        ['Autumn Special',
         '🍂 AUTUMN CAMP — October 2026\n⚽ 7 days in Spain. Professional football development.\n📋 Personal evaluation · 6-month plan · Video analysis\n👨‍🏫 Real Madrid coaching system\n🏨 4★ hotel · All inclusive\n⏰ Early Bird: €2,200 (save €300). Limited spots.',
         'Autumn Camp — Early Bird Open',
         'Sign Up'],
    ],
    col_widths=[2.5, 7, 3, 2]
)

body('8.3 Ретаргетинг реклама (Lead Gen — Неделя 3–8)', bold=True, size=12, color=ACCENT_BLUE)

body('Другой месседж, чем проспектинг. Этот человек уже видел контент / был на сайте.', size=10, color=MEDIUM_GRAY)

table(
    ['Ad', 'Primary Text', 'Headline', 'CTA Button'],
    [
        ['Retarget 1\n"You watched"',
         'You\'ve been exploring football development for your child.\nHere\'s what Voltis Camp gives in 7 days:\n✅ Evaluation against European club benchmarks\n✅ 6-month development plan\n✅ Video analysis of every session\n📍 Spain · All inclusive · Ages 10–19\n🍂 Autumn camp — limited spots.',
         'Ready to Take the Next Step?',
         'Apply Now'],
        ['Retarget 2\nSocial Proof',
         '"My son improved more in 7 days than in 6 months at his local club."\n— Julian\'s dad, Sweden 🇸🇪\n\n📋 Player Report · 📊 10+ metrics · 🎯 6-month plan\n🏆 Real Madrid coaching system\n\nAutumn camp opens soon.',
         'See Why Parents Choose Voltis',
         'Sign Up'],
        ['Retarget 3\nUrgency',
         '⚠️ VOLTIS AUTUMN CAMP — [X] SPOTS LEFT\nOctober 2026 · Spain · 7 days\nReal Madrid coaching. Personal development plan.\nEarly Bird ends [date].\nAfter that — full price or waitlist.',
         'Last Chance — Early Bird Ending',
         'Sign Up'],
    ],
    col_widths=[2.5, 7, 3, 2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 9. CONTENT CALENDAR
# ═══════════════════════════════════════════════════

heading('9. КОНТЕНТ-КАЛЕНДАРЬ НА 8 НЕДЕЛЬ', 1)

body('Расписание публикаций для @voltis.camp. 3–4 поста в неделю + ежедневные Stories.', size=10, color=MEDIUM_GRAY)

weeks = [
    ('НЕДЕЛЯ 1 (1–7 июля) — Запуск', [
        ['Пн', 'POST 1 — Вводный о лагере', 'Single image', 'SELL'],
        ['Вт', 'REEL 1 — "How scouts evaluate"', 'Reel 9:16', 'EDUCATE'],
        ['Чт', 'CAROUSEL 1 — 10 метрик скаутинга', 'Carousel', 'EDUCATE'],
        ['Сб', 'POST 5 — Тренер', 'Single image', 'EDUCATE'],
    ]),
    ('НЕДЕЛЯ 2 (8–14 июля) — Экспертность', [
        ['Пн', 'REEL 2 — "5 mistakes players make"', 'Reel 9:16', 'EDUCATE'],
        ['Ср', 'CAROUSEL 6 — "5 причин почему клуб недостаточно"', 'Carousel', 'EDUCATE'],
        ['Пт', 'REEL 7 — "Real Madrid system — what it means"', 'Reel 9:16', 'EDUCATE'],
        ['Вс', 'POST 7 — Локация (behind the scenes)', 'Single/Carousel', 'SHOW'],
    ]),
    ('НЕДЕЛЯ 3 (15–21 июля) — Social Proof', [
        ['Пн', 'POST 3 — Julian\'s story', 'Single image', 'PROVE'],
        ['Ср', 'REEL 6 — Julian testimonial', 'Reel 9:16', 'PROVE'],
        ['Пт', 'CAROUSEL 7 — Отзывы родителей', 'Carousel', 'PROVE'],
        ['Вс', 'REEL 3 — "Day in the life"', 'Reel 9:16', 'SHOW'],
    ]),
    ('НЕДЕЛЯ 4 (22–31 июля) — Early Bird push', [
        ['Пн', 'POST 4 — Сравнение цен (€357/день)', 'Carousel', 'PROVE'],
        ['Ср', 'CAROUSEL 2 — Voltis vs. Академии', 'Carousel', 'PROVE'],
        ['Пт', 'POST 9 — Early Bird оффер', 'Single image', 'SELL'],
        ['Вс', 'REEL 4 — "€357/day — here\'s what you get"', 'Reel 9:16', 'PROVE'],
    ]),
    ('НЕДЕЛЯ 5 (1–7 августа) — Углубление', [
        ['Пн', 'CAROUSEL 5 — Player Report внутри', 'Carousel', 'PROVE'],
        ['Ср', 'REEL 5 — "What parents don\'t know"', 'Reel 9:16', 'EDUCATE'],
        ['Пт', 'POST 10 — Player Report пример', 'Carousel', 'PROVE'],
        ['Вс', 'REEL 9 — Location reveal', 'Reel 9:16', 'SHOW'],
    ]),
    ('НЕДЕЛЯ 6 (8–14 августа) — Семья + реферал', [
        ['Пн', 'POST 11 — Семейный пакет', 'Single image', 'SELL'],
        ['Ср', 'REEL 8 — "Is your child ready?" quiz', 'Reel 9:16', 'EDUCATE'],
        ['Пт', 'CAROUSEL 3 — 7 дней по шагам', 'Carousel', 'SHOW'],
        ['Вс', 'POST 8 — Возражение "дорого"', 'Text post', 'PROVE'],
    ]),
    ('НЕДЕЛЯ 7 (15–21 августа) — Urgency', [
        ['Пн', 'CAROUSEL 8 — Early Bird: 3 причины', 'Carousel', 'SELL'],
        ['Ср', 'REEL 10 — "This is NOT a summer camp"', 'Reel 9:16', 'SELL'],
        ['Пт', 'POST 6 — "Это не летний лагерь"', 'Single image', 'SELL'],
        ['Вс', 'CAROUSEL 4 — FAQ', 'Carousel', 'SELL'],
    ]),
    ('НЕДЕЛЯ 8 (22–31 августа) — Финал', [
        ['Пн', 'POST 12 — Urgency / last spots', 'Single image', 'SELL'],
        ['Ср', 'REEL 6 (повтор) — Julian testimonial', 'Reel 9:16', 'PROVE'],
        ['Пт', 'POST 9 (обновлённый) — "Price goes up [date]"', 'Single image', 'SELL'],
        ['Вс', 'Итоговый REEL — "See you in October"', 'Reel 9:16', 'SELL'],
    ]),
]

for week_title, posts_list in weeks:
    body(week_title, bold=True, size=11, color=ACCENT_BLUE)
    table(
        ['День', 'Контент', 'Формат', 'Столп'],
        posts_list,
        col_widths=[1.5, 8, 3, 2.5]
    )

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 10. HASHTAGS & VISUAL RULES
# ═══════════════════════════════════════════════════

heading('10. ХЭШТЕГИ И ВИЗУАЛЬНЫЕ ПРАВИЛА', 1)

body('10.1 Набор хэштегов (ротация)', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Группа', 'Хэштеги'],
    [
        ['Основные (всегда)', '#voltiscamp #footballcamp #footballdevelopment #youthfootball'],
        ['Локация', '#footballspain #castellon #mediterranean #spaintraining'],
        ['Продукт', '#playerreport #footballevaluation #developmentplan #footballmetrics'],
        ['Аудитория', '#footballparents #footballmom #footballdad #youngplayer'],
        ['Методология', '#realmadrid #spanishfootball #footballmethodology #footballcoaching'],
        ['Engagement', '#footballtips #footballdrills #footballskills #soccercamp'],
        ['Сезонные', '#autumncamp #octobercamp #earlybird #limitedspots'],
    ],
    col_widths=[3.5, 13]
)

body('Правило: 15–20 хэштегов на пост. Микс из 2–3 групп. Не повторять один набор 2 поста подряд.', size=10, color=MEDIUM_GRAY)

body('10.2 Визуальные правила для ленты', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Правило', 'Описание'],
    [
        ['Сетка 3x3', 'Каждые 9 постов = визуально гармоничный блок. Чередовать: тёмные + яркие + текстовые'],
        ['Не больше 2 продающих постов подряд', 'Между SELL-постами — минимум 1 EDUCATE или SHOW'],
        ['Текст на изображениях', 'Максимум 20% площади — иначе Meta режет охват рекламы'],
        ['Единый фильтр/пресет', 'Все фото — одна цветокоррекция (тёмная/кинематограф ИЛИ яркая/lifestyle)'],
        ['Лого на каждом крео', 'Voltis логотип — левый верхний угол, полупрозрачный'],
        ['CAMP / PROGRAM крупно', 'Требование Adrian — слово CAMP всегда визуально заметно'],
    ],
    col_widths=[4.5, 12]
)

body('10.3 Сводка по объёму контента', bold=True, size=12, color=ACCENT_BLUE)

table(
    ['Тип', 'Количество', 'Статус'],
    [
        ['Reels (готовые скрипты)', '10 шт.', 'Скрипты готовы — нужна съёмка'],
        ['Carousels (раскладки)', '8 шт.', 'Тексты готовы — нужен дизайн в Canva'],
        ['Посты с подписями', '12 шт.', 'Тексты готовы — нужны визуалы'],
        ['Stories-шаблоны', '10 типов', 'Описания готовы — нужны шаблоны в Canva'],
        ['Рекламные тексты (Meta Ads)', '9 шт.', 'Тексты готовы — нужны крео'],
        ['Email-письма', '9 шт.', 'В маркетинговой стратегии'],
        ['ИТОГО единиц контента', '~50+', ''],
    ],
    col_widths=[5.5, 3.5, 6]
)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Конец документа —')
r.font.size = Pt(10); r.font.color.rgb = MEDIUM_GRAY; r.italic = True; r.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Контент-план разработан на основе маркетингового аудита,\nрекламной аналитики и проектной документации Voltis.\nИюнь 2026 · @voltis.camp')
r.font.size = Pt(9); r.font.color.rgb = MEDIUM_GRAY; r.font.name = 'Calibri'

output = '/home/user/tuttofare/Voltis_Content_Plan.docx'
doc.save(output)
print(f'Saved: {output}')
print(f'Size: {os.path.getsize(output) / 1024:.0f} KB')
