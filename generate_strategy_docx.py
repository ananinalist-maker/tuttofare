from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

BRAND_BLACK = RGBColor(0x0d, 0x0d, 0x0d)
BRAND_LIME = RGBColor(0xb5, 0xf0, 0x3c)
ACCENT_GREEN = RGBColor(0x2e, 0x7d, 0x32)
ACCENT_BLUE = RGBColor(0x1a, 0x56, 0x76)
LIGHT_GRAY = RGBColor(0xf5, 0xf5, 0xf5)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xff, 0xff, 0xff)
RED = RGBColor(0xc6, 0x28, 0x28)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLACK
        run.font.name = 'Calibri'
    return h

def add_body(text, bold=False, italic=False, color=None, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(headers, rows, col_widths=None, header_color="1a5676"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f0f0f0")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 70)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(6)

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VOLTIS FOOTBALL\nDEVELOPMENT CAMP')
run.font.size = Pt(36)
run.font.color.rgb = BRAND_BLACK
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('МАРКЕТИНГОВАЯ СТРАТЕГИЯ')
run.font.size = Pt(20)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Дорожная карта на 2 месяца\nЦель: 10 продаж к осеннему лагерю')
run.font.size = Pt(14)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Июнь 2026')
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('КОНФИДЕНЦИАЛЬНО')
run.font.size = Pt(10)
run.font.color.rgb = RED
run.bold = True
run.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════

add_styled_heading('СОДЕРЖАНИЕ', 1)

toc_items = [
    ('1.', 'Резюме проекта'),
    ('2.', 'Текущее состояние: аудит и диагностика'),
    ('3.', 'Цель и обратная математика'),
    ('4.', 'Стратегия: архитектура воронки'),
    ('5.', 'Дорожная карта: 8 недель по шагам'),
    ('6.', 'Декомпозиция задач'),
    ('7.', 'Рекламный бюджет и медиаплан'),
    ('8.', 'Контент-план и материалы'),
    ('9.', 'Система продаж и скрипты'),
    ('10.', 'KPI и контрольные точки'),
    ('11.', 'Риски и план Б'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run_num = p.add_run(f'{num}  ')
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = ACCENT_BLUE
    run_num.font.name = 'Calibri'
    run_t = p.add_run(title)
    run_t.font.size = Pt(12)
    run_t.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════

add_styled_heading('1. РЕЗЮМЕ ПРОЕКТА', 1)

add_body('Продукт', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Параметр', 'Значение'],
    [
        ['Продукт', '7-дневный профессиональный футбольный лагерь'],
        ['Локация', 'Кастельон, Испания'],
        ['Возраст', '10–19 лет'],
        ['Цена', 'от €2,500 (семейный от €1,500)'],
        ['Формат', 'All Inclusive: 4★ отель, 2 тренировки/день, тренер Real Madrid'],
        ['USP', 'Player Report 8–10 стр. + план развития 6 мес. + 30 дней поддержки'],
        ['Instagram', '@voltis.camp'],
        ['Сайт', 'legacy-fc.replit.app'],
    ],
    col_widths=[5, 12]
)

add_body('Задача стратегии', bold=True, size=12, color=ACCENT_BLUE)

p = doc.add_paragraph()
run = p.add_run('Получить 10 оплаченных участников на осенний лагерь (октябрь 2026) за 2 месяца (июль–август), используя платную рекламу Meta Ads, email-маркетинг и прогрев существующей базы лидов.')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_body('Ожидаемый результат', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Показатель', 'Значение'],
    [
        ['Целевые продажи', '10 мест (пессимист 7 / оптимист 15)'],
        ['Целевая выручка', '€22,000–37,500'],
        ['Бюджет на рекламу', '€1,800–2,200'],
        ['Прогнозный ROAS', '10x–17x'],
        ['Горизонт', '8 недель (1 июля — 31 августа)'],
    ],
    col_widths=[5, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 2. AUDIT / CURRENT STATE
# ═══════════════════════════════════════════════════

add_styled_heading('2. ТЕКУЩЕЕ СОСТОЯНИЕ: АУДИТ И ДИАГНОСТИКА', 1)

add_body('2.1 Что уже сделано (итоги мая–июня)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Метрика', 'Значение', 'Оценка'],
    [
        ['Расход на рекламу', '€1,624', ''],
        ['Собрано лидов', '505', 'Отлично'],
        ['Средний CPL', '€3.22', 'Хорошо'],
        ['CPL в мае', '€2.36', 'Отлично'],
        ['CPL в июне', '€5.32', 'Деградация x2.3'],
        ['Качественных лидов', '162 (32%)', 'Норма'],
        ['Продажи', '2 из 505', 'Критично низко'],
        ['Конверсия лид → продажа', '0.4%', 'Проблема'],
        ['Охват', '99,802 чел.', 'Хорошо'],
    ],
    col_widths=[5, 4, 5]
)

add_body('2.2 Главный вывод аудита', bold=True, size=12, color=RED)

p = doc.add_paragraph()
run = p.add_run('Реклама работает. Лиды приходят дёшево. Проблема — в конверсии после лида.')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run('505 лидов → 2 продажи = 0.4% конверсии. При нормальной воронке (5–8%) с теми же лидами было бы 8–13 продаж. Стратегия направлена на решение именно этой проблемы — построить систему прогрева и продаж.')
run.font.size = Pt(10)
run.font.name = 'Calibri'

add_body('2.3 Что работало лучше всего', bold=True, size=12, color=ACCENT_GREEN)

add_table(
    ['Связка', 'CPL', 'Результат'],
    [
        ['LEAD_GENERATION + EU list (NO/SE/BE/NL/DK)', '€2.12–2.74', 'Лучший CPL, 310 лидов'],
        ['Conv-оптимизация + EU list + старая форма', '€7.89 за качеств.', 'Лучшая цена качественного лида'],
        ['Статичные креативы', 'CTR 5.19%', 'В 2 раза выше видео'],
        ['Гео: Скандинавия + Бенилюкс', '€2.53–3.02', 'Лучшее сочетание цена/качество'],
    ],
    col_widths=[7, 3.5, 6]
)

add_body('2.4 Что НЕ работало', bold=True, size=12, color=RED)

add_table(
    ['Проблема', 'Потери', 'Решение'],
    [
        ['50+ адсетов (хаотичная структура)', 'Бюджет размазан, нет обучения', 'Максимум 4 адсета'],
        ['Смешение целей (LEAD_GEN + QUALITY + CONVERSATIONS)', 'Сброс обучения', 'Только LEAD_GENERATION'],
        ['WhatsApp-кампания (CONVERSATIONS)', '€46 = 0 лидов', 'Убрать'],
        ['UK-гео (New Geo)', 'Агенты, мусорные лиды', 'Исключить UK'],
        ['Новая форма с вопросом о бюджете', 'CPL x2.3', 'Простая форма (2 вопроса)'],
        ['Нет follow-up системы', '~500 лидов без обработки', 'Email + CRM + скрипты'],
    ],
    col_widths=[5.5, 4.5, 5.5]
)

add_body('2.5 Что есть и чего не хватает', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Актив', 'Статус', 'Действие'],
    [
        ['Рекламный кабинет Meta', '✓ Активен', 'Реструктуризация'],
        ['Instagram @voltis.camp', '✓ Есть', 'Привязать к рекламному кабинету, если не привязан'],
        ['Facebook-страница', '✓ 79 подписчиков', 'Наполнить контентом'],
        ['Пиксель Meta', '✓ Работает', 'Перенести на нормальный домен'],
        ['База лидов (505 шт.)', '✓ В Google Sheets', 'Загрузить в Custom Audiences'],
        ['WhatsApp Business', '✓ Есть', 'Настроить автоответы'],
        ['4 оффера + 2 видеоскрипта', '✓ Готовы', 'Использовать'],
        ['Email-маркетинг', '✗ Нет', 'Подключить Brevo (бесплатно)'],
        ['Custom Audiences / LAL', '✗ Нет', 'Создать из базы 505 лидов'],
        ['Ретаргетинг', '✗ Нет', 'Запустить на тёплую аудиторию'],
        ['Отзывы / UGC', '✗ Нет', 'Собрать у 2 покупателей'],
        ['Нормальный домен (не Replit)', '✗ Нет', 'Перенести сайт'],
        ['Система обработки лидов', '✗ Нет', 'Скрипты + CRM'],
    ],
    col_widths=[5.5, 3.5, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 3. GOAL & REVERSE MATH
# ═══════════════════════════════════════════════════

add_styled_heading('3. ЦЕЛЬ И ОБРАТНАЯ МАТЕМАТИКА', 1)

add_body('3.1 Почему 10 продаж реально', bold=True, size=12, color=ACCENT_BLUE)

p = doc.add_paragraph()
run = p.add_run('При текущей конверсии 0.4% для 10 продаж нужен бюджет €7,800 — нереалистично. Но конверсия 0.4% — результат отсутствия воронки и системы прогрева. С правильной воронкой конверсия вырастает до 5–8%, и задача решается при бюджете €1,800–2,200.')
run.font.name = 'Calibri'
run.font.size = Pt(10)

add_body('3.2 Расчёт от цели', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Сценарий', 'Конверсия', 'Нужно качеств. лидов', 'Всего лидов', 'Бюджет'],
    [
        ['Текущий (без воронки)', '0.4%', '833', '2,604', '€7,812'],
        ['С прогревом (5%)', '5%', '200', '625', '€1,875'],
        ['С полной воронкой (8%)', '8%', '125', '390', '€1,170'],
    ],
    col_widths=[4.5, 2.5, 4, 3, 3]
)

add_body('3.3 Источники продаж', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Источник', 'Объём', 'Ожид. продажи', 'Логика'],
    [
        ['Прогрев старой базы (162 качеств. лида)', '162 чел.', '3–5', 'Email + WhatsApp + звонки, конверсия 2–3%'],
        ['"Отложенные на осень" (Tracy, Erman и др.)', '~20 чел.', '2–4', 'Уже хотели, ждали осень'],
        ['Новые лиды (реклама июль–август)', '350–450', '2–4', 'CPL €3, 35% качеств., 5% конверсия'],
        ['Рефералы (Julian, Jarno, Tracy)', '3–5 чел.', '1–2', 'Tracy уже привёл семью'],
        ['ИТОГО', '', '7–15', ''],
    ],
    col_widths=[5.5, 2.5, 3, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 4. FUNNEL ARCHITECTURE
# ═══════════════════════════════════════════════════

add_styled_heading('4. СТРАТЕГИЯ: АРХИТЕКТУРА ВОРОНКИ', 1)

add_body('4.1 Текущая модель vs. новая', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['', 'Сейчас', 'Стратегия'],
    [
        ['Путь клиента', 'Реклама → лид-форма → менеджер', '4 уровня: контент → прогрев → оффер → продажа'],
        ['Касаний до продажи', '1–2', '7–12'],
        ['Месседж холодному', '"Запишись на лагерь за €2,500"', '"Узнай как оценивают игроков в Европе"'],
        ['Инструменты', 'Только Meta Ads', 'Meta Ads + Email + WhatsApp + контент'],
        ['Ретаргетинг', 'Нет', 'Video viewers + site visitors + engagement'],
    ],
    col_widths=[4, 6, 6.5]
)

add_body('4.2 Четыре уровня воронки', bold=True, size=12, color=ACCENT_BLUE)

# Level 1
p = doc.add_paragraph()
run = p.add_run('УРОВЕНЬ 1 — ХОЛОДНЫЙ ТРАФИК ')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT_BLUE; run.font.name = 'Calibri'
run2 = p.add_run('(не знают о Voltis)')
run2.font.size = Pt(10); run2.font.color.rgb = MEDIUM_GRAY; run2.font.name = 'Calibri'

add_bullet('Контентная реклама: образовательные видео 45–60 сек (не продающие)')
add_bullet('Органический контент: Instagram @voltis.camp — 3–4 поста в неделю')
add_bullet('Цель: собрать аудиторию Video Viewers 50%+ и IG Engagement для ретаргета')

p = doc.add_paragraph()
run = p.add_run('УРОВЕНЬ 2 — ТЁПЛЫЙ ТРАФИК ')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT_BLUE; run.font.name = 'Calibri'
run2 = p.add_run('(знают, но не готовы)')
run2.font.size = Pt(10); run2.font.color.rgb = MEDIUM_GRAY; run2.font.name = 'Calibri'

add_bullet('Ретаргетинг: тем, кто смотрел видео / был на сайте / взаимодействовал с IG')
add_bullet('Email-цепочка: 5 писем — от образования к офферу (интервал 3 дня)')
add_bullet('WhatsApp: контент + напоминания для тех, кто в диалоге')

p = doc.add_paragraph()
run = p.add_run('УРОВЕНЬ 3 — ГОРЯЧИЙ ТРАФИК ')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT_BLUE; run.font.name = 'Calibri'
run2 = p.add_run('(готовы рассматривать)')
run2.font.size = Pt(10); run2.font.color.rgb = MEDIUM_GRAY; run2.font.name = 'Calibri'

add_bullet('Прямой оффер: "Осенний лагерь — Early Bird €2,200 (экономия €300)"')
add_bullet('Консультация: 15-минутный видеозвонок с менеджером')
add_bullet('Urgency: ограниченные места, дедлайн Early Bird')

p = doc.add_paragraph()
run = p.add_run('УРОВЕНЬ 4 — ПРОДАЖА + ПОСТ-ПРОДАЖА')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = ACCENT_BLUE; run.font.name = 'Calibri'

add_bullet('Оплата: полная или рассрочка 3 платежа')
add_bullet('Подготовка к лагерю: чек-лист, контакт тренера')
add_bullet('Реферал: €200 скидка за приведённую семью')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 5. ROADMAP
# ═══════════════════════════════════════════════════

add_styled_heading('5. ДОРОЖНАЯ КАРТА: 8 НЕДЕЛЬ', 1)

# WEEK 0
add_body('НЕДЕЛЯ 0  |  24–30 июня  |  ФУНДАМЕНТ', bold=True, size=13, color=WHITE)
# Hack: color the paragraph background
# Actually let's use a table for visual effect

week_data = [
    ('НЕДЕЛЯ 0', '24–30 июня', 'ФУНДАМЕНТ', '€0 на рекламу', [
        'Привязать @voltis.camp к рекламному кабинету (если не привязан)',
        'Создать 5 Custom Audiences: Site Visitors, Form Openers, Form Submitters, Video Viewers, IG Engagement',
        'Создать 2 LAL-аудитории (1%) на базе 505 лидов по TIER 1 странам',
        'Подключить Brevo (email-маркетинг) — импорт 162 качественных лидов',
        'Написать 5 писем email-серии (шаблоны см. раздел 8)',
        'Перенести сайт с Replit на нормальный домен (Vercel + voltiscamp.com)',
        'Настроить WhatsApp Business: автоответ + быстрые ответы',
        'Запросить видео-отзывы у Julian (SE) и Jarno (BE)',
        'Согласовать с Adrian: даты осеннего лагеря, цены, Early Bird',
    ]),
    ('НЕДЕЛЯ 1–2', '1–14 июля', 'ПРОГРЕВ', '€210–280', [
        'Запустить email-серию по 162 качественным лидам (5 писем, каждые 3 дня)',
        'WhatsApp-рассылка по "отложенным на осень" лидам (Tracy, Erman и др.)',
        'Прозвонить 6 горячих лидов: Ahmad, Kalina, Kristina, Bjørn, Tracy, Erman',
        'Запустить кампанию CONTENT (Video Views) — €8–10/день, гео TIER 1',
        '3 контентных видео: "Как скауты оценивают", "5 ошибок", "День в лагере"',
        'Органический контент Instagram: 3–4 поста в неделю',
        'Подготовить креативы для ретаргета (на следующий этап)',
    ]),
    ('НЕДЕЛЯ 3–4', '15–31 июля', 'РЕТАРГЕТ + ЛИДГЕН', '€450–550', [
        'Запустить кампанию RETARGET — €8–10/день по тёплой аудитории (video viewers + site visitors)',
        'Запустить кампанию PROSPECT (Lead Gen) — LAL 1% + Broad Scandinavia + Benelux',
        'Новая лид-форма "Autumn Camp" (2 вопроса + ручной номер)',
        'Email #4: оффер Autumn Camp + Early Bird для всей базы',
        'Дедлайн Early Bird Tier 1 (€2,200) — 31 июля',
        'Контрольная точка: сколько продаж? Если 0 — прозвонить ТОП-20 лидов вручную',
    ]),
    ('НЕДЕЛЯ 5–6', '1–14 августа', 'МАСШТАБИРОВАНИЕ', '€500–680', [
        'Масштабирование лучших адсетов: +20% бюджет каждые 3 дня',
        'Худшие адсеты — выключить, бюджет перераспределить',
        'Early Bird Tier 2: €2,350 (экономия €150)',
        'Запустить реферальную программу (€200 скидка за приведённую семью)',
        'Прозвонить ВСЕ лиды в статусе "Дорого" и "Общаемся" с новым оффером',
        'Видеозвонки с горячими лидами (15 мин, снятие возражений)',
    ]),
    ('НЕДЕЛЯ 7–8', '15–31 августа', 'ДОЖИМ И ЗАКРЫТИЕ', '€650–750', [
        'Urgency-кампания: "Осталось X мест", увеличенный бюджет €15–18/день',
        'Email #8: "Last call — Early Bird заканчивается через 48 часов"',
        'Email #9: "Цена вырастет завтра + отзывы родителей"',
        'Финальный прозвон всей базы горячих лидов',
        'Закрытие Early Bird → переход на полную цену €2,500',
        'Итоговый отчёт: X продаж из 10, план на сентябрь',
    ]),
]

for week_title, dates, phase, budget, tasks in week_data:
    # Week header as table
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    cells = t.rows[0].cells
    for i, txt in enumerate([week_title, dates, phase, budget]):
        cells[i].text = ''
        p = cells[i].paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cells[i], "1a5676")

    for task in tasks:
        add_bullet(task)

    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 6. TASK DECOMPOSITION
# ═══════════════════════════════════════════════════

add_styled_heading('6. ДЕКОМПОЗИЦИЯ ЗАДАЧ', 1)

add_body('6.1 Задачи Неделя 0 — Фундамент (до старта рекламы)', bold=True, size=11, color=ACCENT_BLUE)

add_table(
    ['#', 'Задача', 'Ответственный', 'Срок', 'Статус'],
    [
        ['1', 'Привязать @voltis.camp к рекламному кабинету', 'Таргетолог', '24 июня', '☐'],
        ['2', 'Создать 5 Custom Audiences в Meta', 'Таргетолог', '25 июня', '☐'],
        ['3', 'Создать 2 LAL-аудитории (1%) TIER 1', 'Таргетолог', '25 июня', '☐'],
        ['4', 'Зарегистрировать Brevo, импорт 162 лидов', 'Маркетолог', '26 июня', '☐'],
        ['5', 'Написать 5 email-писем (тексты — раздел 8)', 'Копирайтер', '27 июня', '☐'],
        ['6', 'Купить домен voltiscamp.com', 'Разработчик', '25 июня', '☐'],
        ['7', 'Перенести сайт с Replit на Vercel/Netlify', 'Разработчик', '28 июня', '☐'],
        ['8', 'Установить пиксель на новый домен', 'Разработчик', '28 июня', '☐'],
        ['9', 'Настроить WhatsApp: автоответ + быстрые ответы', 'Менеджер', '26 июня', '☐'],
        ['10', 'Запросить отзывы у Julian и Jarno', 'Менеджер', '24 июня', '☐'],
        ['11', 'Согласовать осенние даты и цены с Adrian', 'Руководитель', '25 июня', '☐'],
        ['12', 'Подготовить Early Bird ценовую таблицу', 'Маркетолог', '27 июня', '☐'],
    ],
    col_widths=[0.8, 7, 3, 2.5, 1.5]
)

add_body('6.2 Задачи Неделя 1–2 — Прогрев', bold=True, size=11, color=ACCENT_BLUE)

add_table(
    ['#', 'Задача', 'Ответственный', 'Срок', 'Статус'],
    [
        ['13', 'Запустить email-серию (автоматизация в Brevo)', 'Маркетолог', '1 июля', '☐'],
        ['14', 'WhatsApp-рассылка по "отложенным" лидам', 'Менеджер', '1 июля', '☐'],
        ['15', 'Прозвонить 6 горячих лидов (список в разделе 9)', 'Менеджер', '1–3 июля', '☐'],
        ['16', 'Снять 3 контентных видео (скрипты — раздел 8)', 'Видеограф', '1–5 июля', '☐'],
        ['17', 'Запустить кампанию CONTENT (Video Views)', 'Таргетолог', '3 июля', '☐'],
        ['18', 'Публикация органического контента (8 постов за 2 нед.)', 'SMM', '1–14 июля', '☐'],
        ['19', 'Подготовить 3 креатива для ретаргета', 'Дизайнер', '10 июля', '☐'],
        ['20', 'Анализ open rate / click rate email-серии', 'Маркетолог', '7 июля', '☐'],
    ],
    col_widths=[0.8, 7, 3, 2.5, 1.5]
)

add_body('6.3 Задачи Неделя 3–4 — Ретаргет + лидген', bold=True, size=11, color=ACCENT_BLUE)

add_table(
    ['#', 'Задача', 'Ответственный', 'Срок', 'Статус'],
    [
        ['21', 'Создать лид-форму "Autumn Camp"', 'Таргетолог', '15 июля', '☐'],
        ['22', 'Запустить кампанию RETARGET (тёплая аудитория)', 'Таргетолог', '15 июля', '☐'],
        ['23', 'Запустить кампанию PROSPECT (LAL + Broad)', 'Таргетолог', '15 июля', '☐'],
        ['24', 'Email #4: Autumn Camp оффер → вся база', 'Маркетолог', '18 июля', '☐'],
        ['25', 'Прозвон ТОП-20 лидов (если 0 продаж)', 'Менеджер', '25 июля', '☐'],
        ['26', 'Дедлайн Early Bird Tier 1 (€2,200)', 'Все', '31 июля', '☐'],
        ['27', 'Контрольный отчёт: лиды, продажи, CPL', 'Маркетолог', '31 июля', '☐'],
    ],
    col_widths=[0.8, 7, 3, 2.5, 1.5]
)

add_body('6.4 Задачи Неделя 5–8 — Масштабирование и закрытие', bold=True, size=11, color=ACCENT_BLUE)

add_table(
    ['#', 'Задача', 'Ответственный', 'Срок', 'Статус'],
    [
        ['28', 'Масштабирование лучших адсетов (+20%/3 дня)', 'Таргетолог', '1–14 авг', '☐'],
        ['29', 'Запустить реферальную программу', 'Менеджер', '1 августа', '☐'],
        ['30', 'Прозвон ВСЕХ лидов "Дорого" / "Общаемся"', 'Менеджер', '1–7 авг', '☐'],
        ['31', 'Видеозвонки с горячими (15 мин)', 'Менеджер', '1–14 авг', '☐'],
        ['32', 'Urgency-кампания "X мест осталось"', 'Таргетолог', '15 августа', '☐'],
        ['33', 'Email #8: "Last call" + Email #9: отзывы', 'Маркетолог', '25–28 авг', '☐'],
        ['34', 'Финальный прозвон всей горячей базы', 'Менеджер', '25–30 авг', '☐'],
        ['35', 'Закрытие Early Bird → полная цена', 'Руководитель', '31 августа', '☐'],
        ['36', 'Итоговый отчёт за 2 месяца', 'Маркетолог', '31 августа', '☐'],
    ],
    col_widths=[0.8, 7, 3, 2.5, 1.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 7. BUDGET & MEDIA PLAN
# ═══════════════════════════════════════════════════

add_styled_heading('7. РЕКЛАМНЫЙ БЮДЖЕТ И МЕДИАПЛАН', 1)

add_body('7.1 Структура рекламного кабинета (целевая)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Кампания', 'Цель', 'Адсеты', 'Бюджет/день', 'Период'],
    [
        ['1. CONTENT', 'Video Views', 'Broad EU TIER 1 (30–55 лет)', '€8–10', 'Нед. 1–4'],
        ['2. PROSPECT', 'Lead Generation', 'A: Скандинавия (NO/SE/DK/FI)\nB: Бенилюкс+ (BE/NL/CH/AT)\nC: LAL 1%', '€12–15\n€12–15\n€8–10', 'Нед. 3–8'],
        ['3. RETARGET', 'Lead Generation', 'Video Viewers + Site Visitors + IG Engagement\n(исключить Form Submitters)', '€8–10', 'Нед. 3–8'],
    ],
    col_widths=[2.5, 2.5, 5.5, 2.5, 2]
)

add_body('Всего: 3 кампании, 4 адсета, ~12 объявлений. Дневной бюджет: €40–50', bold=True, size=10)

add_body('7.2 Бюджет по неделям', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Период', 'Контент (VV)', 'Проспектинг (LG)', 'Ретаргет (LG)', 'ИТОГО'],
    [
        ['Нед. 0 (подготовка)', '€0', '€0', '€0', '€0'],
        ['Нед. 1–2 (прогрев)', '€140–200', '€0', '€0', '€140–200'],
        ['Нед. 3–4 (лидген)', '€100–140', '€224–300', '€112–140', '€436–580'],
        ['Нед. 5–6 (масштаб)', '€50–70', '€280–350', '€140–180', '€470–600'],
        ['Нед. 7–8 (дожим)', '€0', '€350–450', '€180–210', '€530–660'],
        ['ИТОГО', '€290–410', '€854–1,100', '€432–530', '€1,576–2,040'],
    ],
    col_widths=[3.5, 3, 3, 3, 3]
)

add_body('7.3 Полный бюджет (реклама + инфраструктура)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Статья', 'Июль', 'Август', 'Итого', 'Примечание'],
    [
        ['Meta Ads — контент (Video Views)', '€200–280', '€50–70', '€250–350', 'Сбор аудитории для ретаргета'],
        ['Meta Ads — проспектинг (Lead Gen)', '€224–300', '€630–800', '€854–1,100', 'LAL + Broad по TIER 1 гео'],
        ['Meta Ads — ретаргетинг (Lead Gen)', '€112–140', '€320–390', '€432–530', 'Тёплая аудитория → лид-форма'],
        ['Итого Meta Ads', '€536–720', '€1,000–1,260', '€1,536–1,980', ''],
        ['Домен (voltiscamp.com)', '€12', '€0', '€12', 'Годовая регистрация'],
        ['Хостинг (Vercel/Netlify)', '€0', '€0', '€0', 'Бесплатный тариф'],
        ['Email-сервис (Brevo)', '€0', '€0', '€0', 'Бесплатно до 300/день'],
        ['WhatsApp Business', '€0', '€0', '€0', 'Бесплатное приложение'],
        ['ОБЩИЙ БЮДЖЕТ', '€548–732', '€1,000–1,260', '€1,548–1,992', ''],
    ],
    col_widths=[5, 2.5, 2.5, 2.5, 4]
)

add_body('7.4 Гео-стратегия (куда показываем рекламу)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Тier', 'Страны', 'CPL (доказанный)', 'Доля бюджета', 'Почему'],
    [
        ['TIER 1', 'NO, SE, DK, BE, NL, CH, AT', '€2.00–3.02', '70%', 'Лучший CPL + платёжеспособность + обе продажи отсюда'],
        ['TIER 2', 'FI, FR, LU', '€3.64–4.26', '15%', 'Средний CPL, тестируем'],
        ['ТЕСТ', 'RO, BG, HR, IE, AL', '€0.47–0.90', '15%', 'Дешёвый CPL, проверяем качество'],
        ['Исключить', 'ES, UK, LT, EE, LV, IT, PT', '—', '0%', 'Мусор / не ЦА / дорого'],
    ],
    col_widths=[2, 5, 3, 2.5, 5]
)

add_body('7.5 Прогноз результатов', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Метрика', 'Пессимист', 'Базовый', 'Оптимист'],
    [
        ['Бюджет на рекламу', '€1,550', '€1,800', '€2,000'],
        ['Новые лиды', '300', '400', '500'],
        ['CPL', '€3.50', '€3.00', '€2.50'],
        ['Качественных лидов (35%)', '105', '140', '175'],
        ['+ Прогрев старой базы', '+20 горячих', '+25 горячих', '+30 горячих'],
        ['Конверсия в продажу', '4%', '6%', '8%'],
        ['ПРОДАЖИ', '7', '10', '15'],
        ['Выручка', '€15,400', '€23,000', '€37,500'],
        ['ROAS', '9.9x', '12.8x', '18.8x'],
    ],
    col_widths=[5, 3.5, 3.5, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 8. CONTENT PLAN
# ═══════════════════════════════════════════════════

add_styled_heading('8. КОНТЕНТ-ПЛАН И МАТЕРИАЛЫ', 1)

add_body('8.1 Видео для рекламы (3 шт. — контентная кампания)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['#', 'Тема', 'Длина', 'Формат', 'Ключевой месседж'],
    [
        ['V1', 'How scouts evaluate players', '60 сек', 'Talking head + графика', '10 метрик оценки — образовательный, НЕ продающий'],
        ['V2', '5 mistakes young players make', '45 сек', 'Talking head на поле', 'Тренер разбирает ошибки — экспертность'],
        ['V3', 'Day in the life at Voltis Camp', '60 сек', 'Lifestyle / B-roll', 'Утро → тренировки → питание → отчёт — атмосфера'],
    ],
    col_widths=[0.8, 4.5, 1.5, 3.5, 6]
)

add_body('8.2 Креативы для ретаргета (3 шт.)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['#', 'Тип', 'Месседж', 'CTA'],
    [
        ['R1', 'Статика 1:1', '"You watched, now act" — для тех, кто видел видео', 'Apply now → Early Bird'],
        ['R2', 'Статика 1:1', 'Social proof — цитата отзыва Julian/Jarno', 'Reserve your spot'],
        ['R3', 'Статика 1:1', 'Urgency — "Only X spots left"', 'Secure your spot NOW'],
    ],
    col_widths=[0.8, 2.5, 8, 4]
)

add_body('8.3 Email-серия (9 писем)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['#', 'День', 'Тема письма', 'Тип', 'CTA'],
    [
        ['0', 'Сразу', 'Welcome — here\'s what happens next', 'Автоматическое', 'Читать статью'],
        ['1', 'День 3', 'The 10 metrics that matter in European football', 'Образование', 'Статья на сайте'],
        ['2', 'День 6', 'Why talent alone won\'t get your child into a club', 'Проблема', 'Видео методологии'],
        ['3', 'День 9', '"My son improved more in 7 days..." — Julian\'s story', 'Social proof', 'Видео-отзыв'],
        ['4', 'День 12', 'Autumn Camp — Early Bird pricing', 'Оффер', 'Забронировать'],
        ['5', 'День 15', 'Your questions answered (FAQ)', 'Возражения', 'WhatsApp'],
        ['6', 'Нед. 5', 'What\'s included — full day-by-day breakdown', 'Детали', 'Забронировать'],
        ['7', 'Нед. 7', 'Last call — Early Bird ends in 48 hours', 'Urgency', 'Забронировать'],
        ['8', 'Нед. 8', 'Price goes up tomorrow + parent reviews', 'Финальный', 'Забронировать'],
    ],
    col_widths=[0.8, 1.5, 6, 2.5, 3]
)

add_body('8.4 Органический контент Instagram (3–4 поста/неделю)', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['День', 'Тип', 'Формат', 'Пример темы'],
    [
        ['Пн', 'Образовательный', 'Carousel', '"5 метрик, по которым скауты оценивают игрока"'],
        ['Ср', 'Social Proof', 'Reel / Story', 'Отзыв Julian + кадры с тренировки'],
        ['Пт', 'Behind the scenes', 'Reel', 'Тренер объясняет упражнение / площадка / отель'],
        ['Вс', 'Engagement', 'Story с опросом', '"В каком возрасте лучше начинать просмотры?"'],
    ],
    col_widths=[1.5, 3, 3, 8]
)

add_body('8.5 Early Bird ценовая механика', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Период', 'Цена', 'Скидка', 'Доступно мест'],
    [
        ['До 31 июля (Tier 1)', '€2,200', '-€300 (12%)', '5'],
        ['1–15 августа (Tier 2)', '€2,350', '-€150 (6%)', '5'],
        ['После 15 августа', '€2,500', 'Полная цена', 'По наличию'],
        ['Семейный пакет Early Bird', '€1,300', '-€200', 'Ограничено'],
    ],
    col_widths=[5, 3, 3, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 9. SALES SYSTEM
# ═══════════════════════════════════════════════════

add_styled_heading('9. СИСТЕМА ПРОДАЖ', 1)

add_body('9.1 Скрипт обработки лидов', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Этап', 'Когда', 'Действие'],
    [
        ['Первый контакт', '< 1 час после лида', 'WhatsApp: приветствие + 1 вопрос о ребёнке'],
        ['Квалификация', '0–2 часа', 'Узнать: возраст, уровень, цели, бюджет'],
        ['Презентация', '2–24 часа', 'PDF-презентация + видео-отзыв'],
        ['Follow-up 1', 'День 2', '"Посмотрели материалы? Какие вопросы?"'],
        ['Видеозвонок', 'День 3–5', '15 мин: ответы + снятие возражений'],
        ['Оффер', 'После звонка', 'Цена + Early Bird + дедлайн + рассрочка'],
        ['Follow-up 2', 'День 7', '"Осталось X мест — есть вопросы?"'],
        ['Follow-up 3', 'День 14', '"Последний шанс по ранней цене"'],
        ['Перевод', 'День 21', 'Если нет ответа → email-nurturing'],
    ],
    col_widths=[3, 3, 10]
)

add_body('9.2 Обработка возражений', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Возражение', 'Ответ'],
    [
        ['"Дорого"', '"€357/день: 2 тренировки с тренером Real Madrid, 4★ отель, all inclusive, план на 6 мес. Годовая программа в академии — €18,000–50,000. Voltis — в 7–20 раз дешевле"'],
        ['"Далеко"', '"All inclusive: перелёт до Валенсии, трансфер, проживание, питание — всё решено. Семейный пакет — вы тоже отдыхаете в Испании"'],
        ['"Ребёнок не готов"', '"Мы берём с 10 лет. Это НЕ просмотр — это тренировка + оценка. Player Report покажет текущий уровень и что делать дальше"'],
        ['"Подумаю"', '"Конечно. Отправлю видео-отзыв семьи из [страна]. Early Bird цена действует до [дата] — экономия €300"'],
        ['"Осенью, не летом"', '"Отлично! Осенний лагерь — октябрь. Early Bird до [дата]. Забронировать?"'],
    ],
    col_widths=[3, 13.5]
)

add_body('9.3 Приоритетные лиды для прозвона', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Лид', 'Страна', 'Статус', 'Действие'],
    [
        ['Ahmad Terkawi', 'Швеция', '2 игрока, реагировал на цены', 'Early Bird на 2 места'],
        ['Kalina Van der Poel', 'Бельгия', 'Сын 13 лет, высокий интерес', 'Звонок + презентация'],
        ['Kristina Palushaj', 'Норвегия', 'Очень заинтересована', 'Звонок + оффер'],
        ['Bjørn Berget', 'Норвегия', 'Семья, сын 10 лет', 'Семейный пакет Early Bird'],
        ['Tracy Sita', 'Бельгия', 'Перенёс на осень + привёл семью', 'VIP-оффер + реферальная скидка'],
        ['Erman', 'Бельгия', '"Лето дорого, осень — да"', 'Early Bird €2,200'],
    ],
    col_widths=[3.5, 2.5, 5, 5.5]
)

add_body('9.4 Реферальная программа', bold=True, size=12, color=ACCENT_BLUE)

p = doc.add_paragraph()
run = p.add_run('Для существующих покупателей (Julian, Jarno) и новых:')
run.font.name = 'Calibri'; run.font.size = Pt(10)

add_bullet('"Пригласите семью — получите €200 скидку на следующий лагерь"')
add_bullet('"Друг тоже получает €100 скидку на первый лагерь"')
add_bullet('Tracy Sita уже привёл семью органически — значит канал работает')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 10. KPIs
# ═══════════════════════════════════════════════════

add_styled_heading('10. KPI И КОНТРОЛЬНЫЕ ТОЧКИ', 1)

add_body('10.1 Целевые KPI', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Метрика', 'Текущее', 'Цель', 'Как достигаем'],
    [
        ['CPL (общий)', '€3.22 → €5.32', '€2.50–3.50', 'Оптимизация гео + LAL + структура'],
        ['CPL качественный', '~€16 (июнь)', '€8–10', 'TIER 1 гео + LAL аудитории'],
        ['CTR', '2.54% (июнь)', '>4.0%', 'Лучшие статики + A/B тест'],
        ['Конверсия лид → диалог', '~32%', '>40%', 'Скорость ответа <1ч + автоответ WhatsApp'],
        ['Конверсия диалог → продажа', '~1.2%', '>5%', 'Воронка + скрипты + follow-up'],
        ['Email Open Rate', '—', '>25%', 'Персонализация + цепляющие темы'],
        ['Email Click Rate', '—', '>5%', 'Релевантный контент + CTA'],
        ['Частота показов', '2.34', '<2.0', 'Не более 2 недель на одном крео'],
    ],
    col_widths=[4, 3, 2.5, 7]
)

add_body('10.2 Еженедельный дашборд', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Метрика', 'Нед 1', 'Нед 2', 'Нед 3', 'Нед 4', 'Нед 5', 'Нед 6', 'Нед 7', 'Нед 8'],
    [
        ['Расход (€)', '', '', '', '', '', '', '', ''],
        ['Лиды (новые)', '', '', '', '', '', '', '', ''],
        ['CPL (€)', '', '', '', '', '', '', '', ''],
        ['Качественных', '', '', '', '', '', '', '', ''],
        ['В диалоге', '', '', '', '', '', '', '', ''],
        ['Видеозвонков', '', '', '', '', '', '', '', ''],
        ['Выставлено счетов', '', '', '', '', '', '', '', ''],
        ['ОПЛАТЫ', '', '', '', '', '', '', '', ''],
        ['Email opens (%)', '', '', '', '', '', '', '', ''],
    ],
    col_widths=[3, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]
)

add_body('10.3 Красные флаги и действия', bold=True, size=12, color=RED)

add_table(
    ['Когда', 'Красный флаг', 'Немедленное действие'],
    [
        ['Неделя 2', 'CPL > €4.00', 'Сузить гео до TIER 1, вернуть крео мая'],
        ['Неделя 3', '0 ответов на email-серию', 'Проверить спам, сменить тему писем'],
        ['Неделя 4', '0 продаж из прогретой базы', 'Прозвонить вручную ТОП-20'],
        ['Неделя 5', 'CTR < 2% на ретаргете', 'Сменить крео, проверить частоту'],
        ['Неделя 6', '< 3 продаж суммарно', 'Снизить Early Bird + добавить рассрочку'],
        ['Неделя 7', '< 5 продаж суммарно', 'Запустить вебинар (план Б)'],
        ['Неделя 8', '< 8 продаж', 'Партнёрства с клубами + агрессивный дожим'],
    ],
    col_widths=[2.5, 5.5, 8.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 11. RISKS & PLAN B
# ═══════════════════════════════════════════════════

add_styled_heading('11. РИСКИ И ПЛАН Б', 1)

add_body('11.1 Основные риски', bold=True, size=12, color=ACCENT_BLUE)

add_table(
    ['Риск', 'Вероятность', 'Влияние', 'Митигация'],
    [
        ['CPL вырастет выше €5', 'Средняя', 'Высокое', 'Вернуть связку мая: LEAD_GEN + EU list + простая форма'],
        ['Старая база не отвечает', 'Средняя', 'Среднее', 'Прозвон + WhatsApp + новый оффер Early Bird'],
        ['Нет видео-контента (видеограф не найден)', 'Средняя', 'Среднее', 'Использовать статичные крео (CTR 5% доказан)'],
        ['Конверсия осталась <2%', 'Низкая', 'Высокое', 'Вебинар + партнёрства + рассрочка'],
        ['Adrian не согласует Early Bird скидки', 'Низкая', 'Среднее', 'Вместо скидки — бонус (доп. день / доп. сессия)'],
        ['Бюджет урезан', 'Средняя', 'Высокое', 'Фокус только на прогрев базы + органик'],
    ],
    col_widths=[4.5, 2.5, 2, 7.5]
)

add_body('11.2 План Б (если к неделе 5 менее 3 продаж)', bold=True, size=12, color=ACCENT_BLUE)

p = doc.add_paragraph()
run = p.add_run('Вариант 1: Бесплатный вебинар')
run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(11)

add_bullet('Тема: "How to get your child into a European football academy"')
add_bullet('Собрать регистрации через рекламу (CPL для вебинара обычно €1–2)')
add_bullet('В конце — оффер Voltis Camp со скидкой для участников')
add_bullet('Конверсия вебинар → продажа: обычно 3–5%')

p = doc.add_paragraph()
run = p.add_run('Вариант 2: Партнёрства с локальными клубами')
run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(11)

add_bullet('Предложить тренерам/клубам в Скандинавии и Бенилюксе комиссию €200 за каждого приведённого')
add_bullet('Отправить 50 писем по футбольным клубам в NO/SE/BE')
add_bullet('Дать пробный тренировочный день для клуба → конверсия в лагерь')

p = doc.add_paragraph()
run = p.add_run('Вариант 3: Рассрочка')
run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(11)

add_bullet('3 платежа по €833 вместо €2,500 единоразово')
add_bullet('Снижает барьер для лидов в категории "Дорого" (их ~34%)')

add_separator()

# Final note
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Конец документа —')
run.font.size = Pt(10)
run.font.color.rgb = MEDIUM_GRAY
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Документ подготовлен на основе аудита рекламного кабинета Meta Ads\n(€1,624, 505 лидов, 5 кампаний, 50+ адсетов) и проектной документации.\nVoltis Football Development Camp · Июнь 2026')
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

# ── Save ──
output_path = '/home/user/tuttofare/Voltis_Marketing_Strategy_2months.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'Size: {os.path.getsize(output_path) / 1024:.0f} KB')
